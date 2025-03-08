# doc_generator.py
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

def create_recipe_docx(title, ingredients, preparation, filename="receta.docx"):
    document = Document()

    # Configurar la fuente predeterminada a sans-serif (Calibri)
    styles = document.styles
    if 'Normal' in styles:
        style = styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(12)

    # Título: grande, y negrita
    p = document.add_paragraph()
    run = p.add_run(title)
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.bold = True


    # Subtítulo: Ingredientes (negrita y tamaño mediano)
    p = document.add_paragraph()
    run = p.add_run("Ingredientes")
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.bold = True

    # Lista de ingredientes (punteo)
    for item in ingredients.splitlines():
        if item.strip():
            p = document.add_paragraph(item, style='List Bullet')
            # Se asegura que la fuente de la lista sea Calibri y tamaño legible
            p.style.font.name = 'Calibri'
            p.style.font.size = Pt(12)

    # Subtítulo: Preparacion (negrita y tamaño mediano)
    p = document.add_paragraph()
    run = p.add_run("Preparacion")
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.bold = True

    # Párrafo de preparación
    p = document.add_paragraph(preparation)
    p.style.font.name = 'Calibri'
    p.style.font.size = Pt(12)

    # Guardar el documento
    document.save(filename)
